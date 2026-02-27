#!/usr/bin/env python3
"""
Generate a .docx manuscript from The Yakverse Chronicles (Vol 3, Art of Gig).

Reads 13 chapter HTML files from docs/, strips web chrome, and outputs a clean
manuscript suitable for Vellum import.

Named paragraph styles (Vellum-targetable):
  BodyText      — all body paragraphs
  ChapterMeta   — "in which..." subtitle line
  ChapterImage  — paragraphs containing embedded images

Output:
  Publishing/manuscripts/artofgig/artofgig_vol3_YYYYMMDD.docx
  Publishing/manuscripts/artofgig/images/  (exported chapter images)

Usage:
  python3 vol3/generate_docx.py
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ── Paths ────────────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent.parent          # artofgig/
DOCS_DIR = ROOT / "docs"
PUBLISHING_DIR = ROOT.parent                            # Publishing/
MANUSCRIPTS_DIR = PUBLISHING_DIR / "manuscripts" / "artofgig"
IMAGES_DIR = MANUSCRIPTS_DIR / "images"

_TODAY = datetime.now().strftime("%Y%m%d")
OUTPUT_DOCX = MANUSCRIPTS_DIR / ("artofgig_vol3_{}.docx".format(_TODAY))

# ── Chapter list (reading order) ─────────────────────────────────────────────

CHAPTERS = [
    ("Prelude",                        "prelude.html"),
    ("The Shadow's Journey",           "the_shadow_s_journey.html"),
    ("Always Be Strategizing",         "always_be_strategizing.html"),
    ("Making it Interesting",          "making_it_interesting.html"),
    ("The Two Shadows of the Hero",    "the_two_shadows_of_the_hero.html"),
    ("Maneuvers vs. Melees",           "maneuvers_vs_melees.html"),
    ("The Twelve Eigenconversations",  "the_twelve_eigenconversations.html"),
    ("The Shtickbox Affair",           "the_shtickbox_affair.html"),
    ("The Medium is the Client",       "the_medium_is_the_client.html"),
    ("And So it Begins",               "and_so_it_begins.html"),
    ("Staying With the Questions",     "staying_with_the_questions.html"),
    ("Infinity Gig",                   "infinity_gig.html"),
    ("Endgame",                        "endgame.html"),
]

# ── Style helpers ─────────────────────────────────────────────────────────────

def _ensure_style(doc, name, based_on="Normal", font_size=None, italic=False,
                  space_before=None, space_after=None):
    """Return a named paragraph style, creating it if absent."""
    styles = doc.styles
    try:
        return styles[name]
    except KeyError:
        pass
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[based_on]
    if font_size:
        style.font.size = Pt(font_size)
    if italic:
        style.font.italic = True
    if space_before is not None:
        style.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        style.paragraph_format.space_after = Pt(space_after)
    return style


def _register_styles(doc):
    """Register all named paragraph styles used by the generator."""
    _ensure_style(doc, "BodyText",    based_on="Normal",  font_size=11,
                  space_before=0, space_after=8)
    _ensure_style(doc, "ChapterMeta", based_on="Normal",  font_size=10,
                  italic=True, space_before=4, space_after=12)
    _ensure_style(doc, "ChapterImage", based_on="Normal", space_before=6,
                  space_after=6)


# ── HTML → docx helpers ───────────────────────────────────────────────────────

def _strip_links(tag):
    """Replace <a> tags with their inner text, keeping formatting."""
    for a in tag.find_all("a"):
        a.unwrap()


def _para_text(tag):
    """Return text content of a tag with <em>/<strong> rendered inline."""
    return tag.get_text()


def _add_formatted_paragraph(doc, tag, style_name):
    """
    Add a paragraph preserving basic inline formatting (em → italic,
    strong → bold, <br> → newline).  <a> tags are unwrapped first.
    """
    _strip_links(tag)
    para = doc.add_paragraph(style=style_name)
    para.paragraph_format.widow_control = True

    def _walk(node, bold=False, italic=False):
        from bs4 import NavigableString, Tag
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = para.add_run(text)
                run.bold = bold
                run.italic = italic
        elif isinstance(node, Tag):
            if node.name == "br":
                run = para.add_run("\n")
            elif node.name in ("em", "i"):
                for child in node.children:
                    _walk(child, bold=bold, italic=True)
            elif node.name in ("strong", "b"):
                for child in node.children:
                    _walk(child, bold=True, italic=italic)
            else:
                for child in node.children:
                    _walk(child, bold=bold, italic=italic)

    for child in tag.children:
        _walk(child)

    return para


def _embed_image(doc, img_path, ch_num, img_index):
    """
    Embed an image into the doc and export a named copy to IMAGES_DIR.
    Returns the paragraph added, or None if image file not found.
    """
    if not img_path.exists():
        print("  WARNING: image not found: {}".format(img_path))
        return None

    # Export named copy
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    dest_name = "ch{:02d}_img{:02d}{}".format(ch_num, img_index, img_path.suffix)
    shutil.copy2(img_path, IMAGES_DIR / dest_name)

    para = doc.add_paragraph(style="ChapterImage")
    run = para.add_run()
    run.add_picture(str(img_path), width=Inches(4.5))
    return para


# ── Per-chapter rendering ─────────────────────────────────────────────────────

def render_chapter(doc, ch_num, title, html_path):
    """Parse a chapter HTML and append its content to doc."""
    print("Chapter {:2d}: {}".format(ch_num, title))

    with html_path.open("r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    # Strip nav, footer, script, style
    for tag in soup.find_all(["nav", "footer", "script", "style"]):
        tag.decompose()
    for tag in soup.find_all(class_="chapter-nav"):
        tag.decompose()

    main = soup.find("main")
    if main is None:
        print("  WARNING: no <main> found in {}".format(html_path.name))
        return

    # Chapter title (Heading 1)
    h1 = main.find("h1", class_="chapter-title")
    if h1:
        heading_text = h1.get_text(strip=True)
        doc.add_heading(heading_text, level=1)
        h1.decompose()
    else:
        # Fall back to passed-in title
        doc.add_heading(title, level=1)

    # Chapter meta ("in which...")
    meta = main.find(class_="chapter-meta")
    if meta:
        meta_text = meta.get_text(strip=True)
        p = doc.add_paragraph(meta_text, style="ChapterMeta")
        meta.decompose()

    # Walk remaining content in order
    img_index = 0
    for child in list(main.children):
        from bs4 import NavigableString, Tag
        if isinstance(child, NavigableString):
            continue
        if not isinstance(child, Tag):
            continue

        tag_name = child.name

        if tag_name == "p":
            # Skip empty paragraphs
            text = child.get_text(strip=True)
            if not text:
                continue
            _add_formatted_paragraph(doc, child, "BodyText")

        elif tag_name in ("figure", "div") and child.find("img"):
            for img_tag in child.find_all("img"):
                src = img_tag.get("src", "")
                # src is relative to docs/
                img_path = DOCS_DIR / src
                img_index += 1
                _embed_image(doc, img_path, ch_num, img_index)

        elif tag_name == "hr":
            # Section break — add a centred asterism paragraph
            p = doc.add_paragraph("* * *", style="BodyText")
            p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        elif tag_name in ("blockquote",):
            text = child.get_text(strip=True)
            if text:
                doc.add_paragraph(text, style="BodyText")

        elif tag_name in ("h2", "h3", "h4"):
            text = child.get_text(strip=True)
            if text:
                level = int(tag_name[1])
                doc.add_heading(text, level=level)

        elif tag_name in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                text = li.get_text(strip=True)
                if text:
                    doc.add_paragraph(text, style="BodyText")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Ensure output directory exists
    MANUSCRIPTS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _register_styles(doc)

    # Book title page (minimal — Vellum will replace with its own design)
    doc.add_heading("The Yakverse Chronicles", level=0)
    doc.add_paragraph("Volume 3 of Art of Gig", style="ChapterMeta")
    doc.add_paragraph("by Venkatesh Rao", style="BodyText")
    doc.add_page_break()

    for ch_num, (title, filename) in enumerate(CHAPTERS, start=1):
        html_path = DOCS_DIR / filename
        if not html_path.exists():
            print("WARNING: {} not found, skipping".format(filename))
            continue
        render_chapter(doc, ch_num, title, html_path)
        # Page break after each chapter (except last)
        if ch_num < len(CHAPTERS):
            doc.add_page_break()

    doc.save(str(OUTPUT_DOCX))
    print("\nSaved: {}".format(OUTPUT_DOCX))
    if IMAGES_DIR.exists():
        n_images = len(list(IMAGES_DIR.iterdir()))
        print("Images exported: {} files in {}".format(n_images, IMAGES_DIR))


if __name__ == "__main__":
    main()
